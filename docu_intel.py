import streamlit as st
import requests
import json
import os
from pathlib import Path
import fitz  # PyMuPDF
from langchain_openai import AzureChatOpenAI
from tenacity import retry, wait_random_exponential, stop_after_attempt
from docx import Document
from docx.shared import Inches
import io
import re
from PIL import Image

# Configuration
GRAPH_TENANT_ID = "4d4343c6-067a-4794-91f3-5cb10073e5b4"
GRAPH_CLIENT_ID = "5ace14db-3235-4cd2-acfd-dd5ef19d6ea1"
GRAPH_CLIENT_SECRET = "HRk8Q~7G6EH3.yhDC3rB5wLAyAixQMnQNWNyUdsW"
PDF_SITE_ID = "marketingai.sharepoint.com,b82dbaac-09cc-4539-ad08-e4ca926796e8,7b756d20-3463-44b7-95ca-5873f8c3f517"
FUNCTION_URL = "https://doc2pdf.azurewebsites.net"

# Azure OpenAI API details
azure_endpoint = 'https://chat-gpt-a1.openai.azure.com/'
azure_deployment_name = 'DanielChatGPT16k'
azure_api_key = 'c09f91126e51468d88f57cb83a63ee36'
azure_api_version = '2024-05-01-preview'

# Initialize Azure OpenAI LLM
llm = AzureChatOpenAI(
    openai_api_key=azure_api_key,
    api_version=azure_api_version,
    azure_endpoint=azure_endpoint,
    model="gpt-4",
    azure_deployment=azure_deployment_name,
    temperature=0.5
)

def get_oauth2_token():
    url = f"https://login.microsoftonline.com/{GRAPH_TENANT_ID}/oauth2/v2.0/token"
    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    data = {
        'grant_type': 'client_credentials',
        'client_id': GRAPH_CLIENT_ID,
        'client_secret': GRAPH_CLIENT_SECRET,
        'scope': 'https://graph.microsoft.com/.default'
    }
    response = requests.post(url, headers=headers, data=data)
    if response.status_code == 200:
        return response.json().get('access_token')
    else:
        st.error(f"Failed to obtain OAuth2 token: {response.status_code} {response.text}")
        return None

def upload_file_to_sharepoint(token, file):
    upload_url = f"https://graph.microsoft.com/v1.0/sites/{PDF_SITE_ID}/drive/root:/{file.name}:/content"
    headers = {
        'Authorization': f'Bearer {token}',
        'Content-Type': file.type
    }
    response = requests.put(upload_url, headers=headers, data=file.getvalue())
    
    if response.status_code in [200, 201]:
        return response.json().get('id')
    else:
        st.error(f"Failed to upload file to SharePoint: {response.status_code} {response.text}")
        return None

def convert_file_to_pdf(token, file_id):
    convert_url = f"https://graph.microsoft.com/v1.0/sites/{PDF_SITE_ID}/drive/items/{file_id}/content?format=pdf"
    headers = {
        'Authorization': f'Bearer {token}',
        'Content-Type': 'application/json'
    }
    response = requests.get(convert_url, headers=headers)
    if response.status_code == 200:
        return response.content
    else:
        st.error(f"Failed to convert file to PDF: {response.status_code} {response.text}")
        return None

def delete_file_from_sharepoint(token, file_id):
    delete_url = f"https://graph.microsoft.com/v1.0/sites/{PDF_SITE_ID}/drive/items/{file_id}"
    headers = {'Authorization': f'Bearer {token}'}
    response = requests.delete(delete_url, headers=headers)
    if response.status_code == 204:
        return True
    else:
        st.error(f"Failed to delete file from SharePoint: {response.status_code} {response.text}")
        return False

def extract_text_from_pdf(pdf_path: str):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num, page in enumerate(doc):
        page_text = page.get_text()
        st.write(f"Extracting text from Page {page_num + 1}")  # Inform the user about the progress
        text += f"Page {page_num + 1}\n" + page_text
    return text

def extract_images_from_pdf(pdf_path: str, output_folder: str):
    doc = fitz.open(pdf_path)
    images = []
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for page_number in range(len(doc)):
        page = doc[page_number]
        image_list = page.get_images(full=True)
        
        for image_index, img in enumerate(image_list, start=1):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_path = os.path.join(output_folder, f"image_{page_number + 1}_{image_index}.{image_ext}")
            
            with open(image_path, "wb") as img_file:
                img_file.write(image_bytes)
            
            # Simulate meaningful and relevant image titles and descriptions
            image_title = f"Figure {page_number + 1}.{image_index}: Description of key feature"
            image_description = f"This figure illustrates key aspect {image_index} found on page {page_number + 1}."

            images.append({
                "page_number": page_number + 1,
                "title": image_title,
                "description": image_description,
                "image_url": image_path
            })
    
    return images

@retry(wait=wait_random_exponential(min=1, max=120), stop=stop_after_attempt(10))
def completion_with_backoff(prompt: str, content: str):
    try:
        response = llm(
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": content}
            ]
        )
        return response
    except Exception as e:
        st.error(f"Error calling Azure OpenAI API: {e}")
        return None

def extract_metadata(content: str):
    prompt = """Patent Document Analysis and Formatting Prompt:

    You are an expert tasked with analyzing and formatting patent documents. Please thoroughly review the provided patent document and extract the following key information from each page, ensuring that the content is neatly formatted with proper titles, bullet points, and subpoints, maintaining the original structure and spacing:

    1. Page Number: Identify the page number of the document.
    2. Page Title: Extract the title or heading of the page.
    3. Page Content: Extract the exact main content and context of the page as it appears in the document, and format it with bullet points and subpoints. Ensure that each topic and subtopic is clearly identified and described in a structured manner, maintaining the original spacing and indentation. Make the topics and subtopics bold.
    4. Table Content: Identify and extract any tables present on the page, ensuring that the table content is properly structured and formatted.
    5. Image: Identify and extract any images present on the page along with relevant metadata. This includes the image title, description, and any other pertinent information.

    Guidelines:
    - Ensure that all extracted information is factual, accurate, and directly derived from the document.
    - For the "Page Title" and "Image" sections, provide concise and descriptive information.
    - The information should be self-contained, meaning that each extracted piece should make sense independently of the rest of the document.
    - For tables, ensure that the content is properly structured and formatted.
    - For images, include detailed metadata such as:
      - Image Title: The title or caption associated with the image.
      - Image Description: A brief description of the image’s content and purpose.
      - Additional Metadata: Any other relevant details, such as image source or reference numbers.

    Formatting Instructions:
    - Use bullet points for listing items.
    - Use subpoints for detailed explanations under each main point.
    - Ensure that each topic and subtopic is clearly identified and described.
    - Make topics and subtopics bold.
    - Maintain the original structure, spacing, and indentation of the document.
    - Ensure that the content is neat and structured throughout the document.

    Response Format:
    Answer in JSON format. Each page should be represented as an object with the following keys:

    - "PageNumber": The number of the page.
    - "PageTitle": The title of the page as it appears in the document.
    - "PageContent": The exact main content and context of the page as it appears in the document, formatted with bullet points and subpoints, maintaining the original structure and spacing.
    - "Tables": A list of objects containing:
      - "TableTitle": "Title of the Table",
      - "TableContent": "Structured content of the table."
    - "Images": A list of objects containing:
      - "ImageTitle": "Title of the Image",
      - "ImageDescription": "Description of the Image",
      - "AdditionalMetadata": "Any additional metadata about the image."

    Ensure that the JSON output is well-formatted and accurate.
    """
    return completion_with_backoff(prompt, content)

def create_word_file(content, images, output_path: str):
    doc = Document()

    for page in content:
        doc.add_heading(f'Page {page["PageNumber"]}: {page["PageTitle"]}', level=1)
        doc.add_paragraph(page["PageContent"])

        if "Tables" in page and page["Tables"]:
            for table in page["Tables"]:
                doc.add_heading(table["TableTitle"], level=2)
                doc.add_paragraph(table["TableContent"])

        if "Images" in page and page["Images"]:
            for img in page["Images"]:
                doc.add_heading(img["ImageTitle"], level=2)
                doc.add_paragraph(img["ImageDescription"])
    
    # Add images to the document
    for image_info in images:
        doc.add_heading(image_info["title"], level=2)
        doc.add_paragraph(image_info["description"])
        image_path = image_info["image_url"]
        if Path(image_path).exists():
            doc.add_picture(image_path, width=Inches(5.0))
        else:
            st.warning(f"Image not found: {image_path}")

    doc.save(output_path)

def sanitize_text(text):
    return re.sub(r'[^\x00-\x7F]+', ' ', text)  # Replacing non-ASCII characters with a space

# Streamlit App Code
def main():
    st.title("Patent Document Processor")

    uploaded_file = st.file_uploader("Upload a PowerPoint or PDF file", type=["pptx", "pdf"])
    
    if uploaded_file:
        token = get_oauth2_token()
        if not token:
            st.error("Failed to authenticate. Please check your credentials.")
            return
        
        file_type = Path(uploaded_file.name).suffix.lower()
        
        if file_type == '.pptx' or file_type == '.pdf':
            file_id = upload_file_to_sharepoint(token, uploaded_file)
            if file_id:
                st.write("File uploaded successfully. Converting to PDF...")

                pdf_content = convert_file_to_pdf(token, file_id)
                if pdf_content:
                    pdf_path = os.path.join("temp", f"{Path(uploaded_file.name).stem}.pdf")
                    with open(pdf_path, 'wb') as pdf_file:
                        pdf_file.write(pdf_content)
                    
                    st.write("PDF conversion completed.")

                    text = extract_text_from_pdf(pdf_path)
                    sanitized_text = sanitize_text(text)
                    st.write("Text extraction completed.")

                    images = extract_images_from_pdf(pdf_path, "extracted_images")
                    st.write("Image extraction completed.")

                    metadata = extract_metadata(sanitized_text)
                    st.write("Metadata extraction completed.")
                    
                    output_path = os.path.join("output", f"{Path(uploaded_file.name).stem}.docx")
                    create_word_file(metadata, images, output_path)
                    st.write("Word document created successfully.")

                    st.download_button(label="Download Word Document", data=open(output_path, "rb").read(), file_name=output_path)

                    delete_file_from_sharepoint(token, file_id)
                else:
                    st.error("PDF conversion failed.")
            else:
                st.error("Failed to upload file to SharePoint.")
        else:
            st.error("Unsupported file format. Please upload a PowerPoint or PDF file.")

if __name__ == "__main__":
    main()
